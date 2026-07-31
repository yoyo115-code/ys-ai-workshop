from __future__ import annotations

import html
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from app.schemas.resume_export import ResumeEntry, StructuredResume

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Mm, Pt, RGBColor

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import A4, LETTER
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (
        HRFlowable,
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False


CJK_RE = re.compile(r"[\u3400-\u9fff\uf900-\ufaff]")
CJK_FONT_CANDIDATES = (
    Path("/Library/Fonts/Arial Unicode.ttf"),
    Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
    Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
    Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
)
DOCX_CJK_FONT = "Hiragino Sans GB"
DOCX_METADATA_TIMESTAMP = datetime(2000, 1, 1, tzinfo=timezone.utc)
DOCX_ARCHIVE_TIMESTAMP = (2000, 1, 1, 0, 0, 0)


class DocumentRenderError(Exception):
    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code
        self.message = message


class ResumeDocumentRenderer:
    def render(
        self,
        resume: StructuredResume,
        template_key: str,
        export_format: str,
        paper_size: str,
        language: str,
        output_path: Path,
    ) -> None:
        if export_format == "docx":
            self.render_docx(resume, template_key, paper_size, language, output_path)
            return
        if export_format == "pdf":
            self.render_pdf(resume, template_key, paper_size, language, output_path)
            return
        raise DocumentRenderError("unsupported_format", "Unsupported resume export format")

    def render_docx(
        self,
        resume: StructuredResume,
        template_key: str,
        paper_size: str,
        language: str,
        output_path: Path,
    ) -> None:
        if not DOCX_AVAILABLE:
            raise DocumentRenderError(
                "docx_renderer_unavailable",
                "DOCX renderer is unavailable; install python-docx",
            )
        document = Document()
        document.core_properties.author = ""
        document.core_properties.last_modified_by = ""
        document.core_properties.created = DOCX_METADATA_TIMESTAMP
        document.core_properties.modified = DOCX_METADATA_TIMESTAMP
        document.core_properties.revision = 1
        body_font = DOCX_CJK_FONT if CJK_RE.search(resume.model_dump_json()) else "Arial"
        self._configure_docx(document, template_key, paper_size, body_font)
        self._docx_header(document, resume, template_key)
        self._docx_sections(document, resume, template_key, language)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        self._normalize_docx_archive(output_path)

    @staticmethod
    def _normalize_docx_archive(output_path: Path) -> None:
        with zipfile.ZipFile(output_path, "r") as source:
            entries = [(info, source.read(info.filename)) for info in source.infolist()]
        with zipfile.ZipFile(output_path, "w") as target:
            for info, content in entries:
                info.date_time = DOCX_ARCHIVE_TIMESTAMP
                info.extra = b""
                target.writestr(info, content)

    def render_pdf(
        self,
        resume: StructuredResume,
        template_key: str,
        paper_size: str,
        language: str,
        output_path: Path,
    ) -> None:
        if not REPORTLAB_AVAILABLE:
            raise DocumentRenderError(
                "pdf_renderer_unavailable",
                "PDF renderer is unavailable; install reportlab",
            )
        font_name, bold_font = self._pdf_fonts(resume)
        page = A4 if paper_size == "a4" else LETTER
        professional = template_key == "professional"
        margin = 0.68 * inch if professional else 0.75 * inch
        document = SimpleDocTemplate(
            str(output_path),
            pagesize=page,
            rightMargin=margin,
            leftMargin=margin,
            topMargin=margin,
            bottomMargin=margin,
            title=resume.basics.name or "Resume",
            author="",
            subject="Resume export",
        )
        styles = self._pdf_styles(font_name, bold_font, professional)
        story: list = []
        self._pdf_header(story, resume, styles, professional)
        self._pdf_sections(story, resume, styles, language)
        if not story:
            raise DocumentRenderError("empty_resume", "Resume has no renderable content")
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.build(story)

    @staticmethod
    def _configure_docx(
        document, template_key: str, paper_size: str, body_font: str
    ) -> None:
        professional = template_key == "professional"
        section = document.sections[0]
        if paper_size == "a4":
            section.page_width = Mm(210)
            section.page_height = Mm(297)
        else:
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        margin = Inches(0.68 if professional else 0.75)
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)

        normal = document.styles["Normal"]
        normal.font.name = body_font
        normal.font.size = Pt(10 if professional else 10.25)
        normal.paragraph_format.space_after = Pt(3)
        normal.paragraph_format.line_spacing = 1.08 if professional else 1.05
        ResumeDocumentRenderer._set_style_fonts(normal, body_font)

        if "Resume Section" not in document.styles:
            heading = document.styles.add_style("Resume Section", WD_STYLE_TYPE.PARAGRAPH)
        else:
            heading = document.styles["Resume Section"]
        heading.font.name = body_font
        heading.font.size = Pt(11.5 if professional else 11)
        heading.font.bold = True
        heading.font.color.rgb = RGBColor(31, 78, 121) if professional else RGBColor(0, 0, 0)
        heading.paragraph_format.space_before = Pt(9)
        heading.paragraph_format.space_after = Pt(4)
        heading.paragraph_format.keep_with_next = True
        ResumeDocumentRenderer._set_style_fonts(heading, body_font)

        bullet = document.styles["List Bullet"]
        bullet.font.name = body_font
        bullet.font.size = Pt(9.8 if professional else 10)
        bullet.paragraph_format.left_indent = Inches(0.22)
        bullet.paragraph_format.first_line_indent = Inches(-0.15)
        bullet.paragraph_format.space_after = Pt(2)
        bullet.paragraph_format.line_spacing = 1.05
        ResumeDocumentRenderer._set_style_fonts(bullet, body_font)

    @staticmethod
    def _set_style_fonts(style, name: str) -> None:
        style._element.rPr.rFonts.set(qn("w:ascii"), name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_CJK_FONT)

    @staticmethod
    def _set_run_font(run, size: float, bold: bool = False, color: RGBColor | None = None) -> None:
        font_name = DOCX_CJK_FONT if CJK_RE.search(run.text) else "Arial"
        run.font.name = font_name
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font_name)
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font_name)
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)
        run.font.size = Pt(size)
        run.bold = bold
        if color is not None:
            run.font.color.rgb = color

    def _docx_header(self, document, resume: StructuredResume, template_key: str) -> None:
        professional = template_key == "professional"
        if resume.basics.name:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(2)
            self._set_run_font(
                paragraph.add_run(resume.basics.name),
                21 if professional else 19,
                bold=True,
                color=RGBColor(20, 52, 84) if professional else RGBColor(0, 0, 0),
            )
        contacts = self._contacts(resume)
        if contacts:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(6)
            self._set_run_font(
                paragraph.add_run(" | ".join(contacts)),
                9,
                color=RGBColor(70, 82, 96),
            )
            if professional:
                self._paragraph_bottom_border(paragraph, "7C93AD")

    def _docx_sections(
        self, document, resume: StructuredResume, template_key: str, language: str
    ) -> None:
        if resume.basics.summary:
            self._docx_heading(document, self._label("summary", language), template_key)
            self._docx_text(document, resume.basics.summary)
        for key, entries in (
            ("experience", resume.experience),
            ("projects", resume.projects),
            ("education", resume.education),
        ):
            if entries:
                self._docx_heading(document, self._label(key, language), template_key)
                for entry in entries:
                    self._docx_entry(document, entry)
        for key, items in (
            ("skills", resume.skills),
            ("certifications", resume.certifications),
            ("awards", resume.awards),
            ("additional_information", resume.additional_information),
        ):
            if items:
                self._docx_heading(document, self._label(key, language), template_key)
                if key == "skills":
                    self._docx_text(document, " • ".join(items))
                else:
                    for item in items:
                        self._docx_bullet(document, item)

    def _docx_heading(self, document, text: str, template_key: str) -> None:
        paragraph = document.add_paragraph(style="Resume Section")
        paragraph.add_run(text)
        if template_key == "professional":
            self._paragraph_bottom_border(paragraph, "B9C7D8")

    def _docx_text(self, document, text: str) -> None:
        for line in [line for line in text.splitlines() if line.strip()]:
            paragraph = document.add_paragraph(line.strip())
            paragraph.paragraph_format.widow_control = True

    def _docx_entry(self, document, entry: ResumeEntry) -> None:
        if not any((entry.title, entry.organization, entry.location, entry.start_date, entry.end_date, entry.bullet_points)):
            return
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.keep_with_next = bool(entry.bullet_points)
        width = document.sections[0].page_width - document.sections[0].left_margin - document.sections[0].right_margin
        paragraph.paragraph_format.tab_stops.add_tab_stop(width, WD_TAB_ALIGNMENT.RIGHT)
        label_parts = [value for value in (entry.title, entry.organization) if value]
        self._set_run_font(paragraph.add_run(" — ".join(label_parts)), 10, bold=True)
        date = self._date_range(entry)
        meta = " | ".join(value for value in (entry.location, date) if value)
        if meta:
            self._set_run_font(paragraph.add_run(f"\t{meta}"), 9, color=RGBColor(82, 94, 108))
        for bullet in entry.bullet_points:
            self._docx_bullet(document, bullet)

    @staticmethod
    def _docx_bullet(document, text: str) -> None:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(text)
        paragraph.paragraph_format.widow_control = True

    @staticmethod
    def _paragraph_bottom_border(paragraph, color: str) -> None:
        properties = paragraph._p.get_or_add_pPr()
        borders = properties.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            properties.append(borders)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), color)
        borders.append(bottom)

    def _pdf_fonts(self, resume: StructuredResume) -> tuple[str, str]:
        content = resume.model_dump_json()
        if not CJK_RE.search(content):
            return "Helvetica", "Helvetica-Bold"
        for candidate in CJK_FONT_CANDIDATES:
            if not candidate.exists():
                continue
            try:
                if "YSResumeUnicode" not in pdfmetrics.getRegisteredFontNames():
                    pdfmetrics.registerFont(TTFont("YSResumeUnicode", str(candidate)))
                return "YSResumeUnicode", "YSResumeUnicode"
            except Exception:
                continue
        raise DocumentRenderError(
            "pdf_font_unavailable",
            "No Unicode CJK font is available for PDF generation",
        )

    @staticmethod
    def _pdf_styles(font_name: str, bold_font: str, professional: bool) -> dict[str, ParagraphStyle]:
        heading_color = colors.HexColor("#1F4E79") if professional else colors.black
        return {
            "name": ParagraphStyle(
                "ResumeName",
                fontName=bold_font,
                fontSize=21 if professional else 19,
                leading=24,
                alignment=TA_CENTER,
                textColor=colors.HexColor("#143454") if professional else colors.black,
                spaceAfter=3,
            ),
            "contact": ParagraphStyle(
                "ResumeContact",
                fontName=font_name,
                fontSize=8.8,
                leading=11,
                alignment=TA_CENTER,
                textColor=colors.HexColor("#465260"),
                spaceAfter=7,
            ),
            "section": ParagraphStyle(
                "ResumeSection",
                fontName=bold_font,
                fontSize=11.5 if professional else 11,
                leading=14,
                textColor=heading_color,
                spaceBefore=8,
                spaceAfter=3,
                keepWithNext=True,
            ),
            "body": ParagraphStyle(
                "ResumeBody",
                fontName=font_name,
                fontSize=9.7 if professional else 10,
                leading=12.2,
                alignment=TA_LEFT,
                textColor=colors.black,
                spaceAfter=3,
            ),
            "entry": ParagraphStyle(
                "ResumeEntry",
                fontName=font_name,
                fontSize=9.5,
                leading=12,
                textColor=colors.black,
                spaceBefore=3,
                spaceAfter=2,
                keepWithNext=True,
            ),
            "bullet": ParagraphStyle(
                "ResumeBullet",
                fontName=font_name,
                fontSize=9.5 if professional else 9.8,
                leading=12,
                leftIndent=0.18 * inch,
                firstLineIndent=0,
                spaceAfter=2,
            ),
        }

    def _pdf_header(self, story: list, resume: StructuredResume, styles: dict, professional: bool) -> None:
        if resume.basics.name:
            story.append(Paragraph(self._escape(resume.basics.name), styles["name"]))
        contacts = self._contacts(resume)
        if contacts:
            story.append(Paragraph(self._escape(" | ".join(contacts)), styles["contact"]))
        if professional and (resume.basics.name or contacts):
            story.append(HRFlowable(width="100%", thickness=0.6, color=colors.HexColor("#7C93AD"), spaceAfter=4))

    def _pdf_sections(self, story: list, resume: StructuredResume, styles: dict, language: str) -> None:
        if resume.basics.summary:
            self._pdf_heading(story, self._label("summary", language), styles)
            story.append(Paragraph(self._escape(resume.basics.summary).replace("\n", "<br/>"), styles["body"]))
        for key, entries in (
            ("experience", resume.experience),
            ("projects", resume.projects),
            ("education", resume.education),
        ):
            if entries:
                self._pdf_heading(story, self._label(key, language), styles)
                for entry in entries:
                    self._pdf_entry(story, entry, styles)
        for key, items in (
            ("skills", resume.skills),
            ("certifications", resume.certifications),
            ("awards", resume.awards),
            ("additional_information", resume.additional_information),
        ):
            if not items:
                continue
            self._pdf_heading(story, self._label(key, language), styles)
            if key == "skills":
                story.append(Paragraph(self._escape(" • ".join(items)), styles["body"]))
            else:
                story.append(self._pdf_bullets(items, styles))

    @staticmethod
    def _pdf_heading(story: list, text: str, styles: dict) -> None:
        story.append(Paragraph(html.escape(text), styles["section"]))

    def _pdf_entry(self, story: list, entry: ResumeEntry, styles: dict) -> None:
        if not any((entry.title, entry.organization, entry.location, entry.start_date, entry.end_date, entry.bullet_points)):
            return
        label = " — ".join(value for value in (entry.title, entry.organization) if value)
        meta = " | ".join(value for value in (entry.location, self._date_range(entry)) if value)
        line = f"<b>{self._escape(label)}</b>"
        if meta:
            line += f" &nbsp; {self._escape(meta)}"
        story.append(Paragraph(line, styles["entry"]))
        if entry.bullet_points:
            story.append(self._pdf_bullets(entry.bullet_points, styles))

    @staticmethod
    def _pdf_bullets(items: Iterable[str], styles: dict):
        return ListFlowable(
            [ListItem(Paragraph(html.escape(item), styles["bullet"]), leftIndent=0) for item in items],
            bulletType="bullet",
            start="circle",
            leftIndent=0.18 * inch,
            bulletFontName=styles["bullet"].fontName,
            bulletFontSize=6,
            spaceAfter=2,
        )

    @staticmethod
    def _contacts(resume: StructuredResume) -> list[str]:
        basics = resume.basics
        return [value for value in (basics.email, basics.phone, basics.location, *basics.links) if value]

    @staticmethod
    def _date_range(entry: ResumeEntry) -> str:
        if entry.start_date and entry.end_date:
            return f"{entry.start_date} - {entry.end_date}"
        return entry.start_date or entry.end_date

    @staticmethod
    def _escape(value: str) -> str:
        return html.escape(value, quote=True)

    @staticmethod
    def _label(key: str, language: str) -> str:
        labels = {
            "summary": ("个人简介", "SUMMARY"),
            "experience": ("工作经历", "EXPERIENCE"),
            "projects": ("项目经历", "PROJECTS"),
            "education": ("教育经历", "EDUCATION"),
            "skills": ("技能", "SKILLS"),
            "certifications": ("证书", "CERTIFICATIONS"),
            "awards": ("荣誉奖项", "AWARDS"),
            "additional_information": ("补充信息", "ADDITIONAL INFORMATION"),
        }
        zh, en = labels[key]
        if language == "zh":
            return zh
        if language == "en":
            return en
        return f"{zh} / {en}"
