import io
import json
import secrets
import sqlite3
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfReader


PROJECT_ROOT = Path(__file__).resolve().parents[2]
BACKEND_ROOT = PROJECT_ROOT / "backend"
sys.path.insert(0, str(BACKEND_ROOT))

from app.core.config import Settings  # noqa: E402
from app.main import create_app  # noqa: E402
from app.services.document_rendering import DocumentRenderError  # noqa: E402
from app.services.resume_export import ResumeExportService  # noqa: E402


RESUME_TEXT = """Zheng Jieyao
zheng@example.com | +86 138 0000 0000

SUMMARY
Data analyst focused on reliable decision support.

EXPERIENCE
Data Analyst
Example Co
2024 - Present
- Built Python reporting workflows.
- Improved data quality checks.

PROJECTS
Career Workshop
Personal Project
2025 - Present
- Built explainable resume workflows.

EDUCATION
BSc Data Science
Example University
2020 - 2024

SKILLS
Python, SQL, FastAPI

CERTIFICATIONS
Cloud Fundamentals
"""
JOB_DESCRIPTION = "Requires Python, SQL and data analysis experience."


class ExportMockProvider:
    def generate(self, prompt: str, provider: str) -> str:
        if "SECURITY AND EVIDENCE RULES" in prompt:
            return json.dumps(
                {
                    "overall_alignment": "partial_alignment",
                    "covered_requirements": [],
                    "partially_covered_requirements": [],
                    "missing_requirements": [],
                    "uncertain_requirements": [],
                    "resume_expression_issues": [],
                    "qualification_risks": [],
                    "summary": "Synthetic regression result.",
                    "analysis_limitations": ["Synthetic test input only."],
                }
            )
        return f"mock:{provider}:{len(prompt)}"


class ResumeExportTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary_directory = tempfile.TemporaryDirectory()
        self.root = Path(self.temporary_directory.name)
        self.database_path = self.root / "export-test.db"
        self.export_dir = self.root / "private-exports"
        self.settings = Settings(
            database_url=f"sqlite:///{self.database_path}",
            frontend_dir=PROJECT_ROOT / "frontend",
            schema_path=PROJECT_ROOT / "database" / "schema.sql",
            resume_export_dir=self.export_dir,
        )
        self.provider = ExportMockProvider()
        self.client_context = TestClient(create_app(self.settings, self.provider))
        self.client = self.client_context.__enter__()
        self.password = f"ExportTest{secrets.randbelow(100000):05d}"

    def tearDown(self) -> None:
        self.client_context.__exit__(None, None, None)
        self.temporary_directory.cleanup()

    @staticmethod
    def headers(token: str, json_content: bool = False) -> dict[str, str]:
        headers = {"X-Session-Token": token}
        if json_content:
            headers["Content-Type"] = "application/json"
        return headers

    def register(self, username: str = "export_user") -> dict:
        response = self.client.post(
            "/auth/register",
            json={
                "username": username,
                "password": self.password,
                "display_name": "Export User",
            },
        )
        self.assertEqual(response.status_code, 201, response.text)
        return response.json()

    def workspace(
        self,
        token: str,
        resume_text: str = RESUME_TEXT,
        company_name: str = "Microsoft",
        job_title: str = "Data Analyst",
    ) -> tuple[dict, dict]:
        created = self.client.post(
            "/career/applications",
            headers=self.headers(token),
            data={
                "resume_text": resume_text,
                "job_description": JOB_DESCRIPTION,
                "company_name": company_name,
                "job_title": job_title,
                "language": "bilingual",
            },
        )
        self.assertEqual(created.status_code, 201, created.text)
        application = created.json()
        opened = self.client.get(
            f"/career/applications/{application['id']}/resume-suggestions",
            headers=self.headers(token),
        )
        self.assertEqual(opened.status_code, 200, opened.text)
        return application, opened.json()

    def preview(self, token: str, version_id: int) -> dict:
        response = self.client.get(
            f"/career/resume-versions/{version_id}/preview",
            headers=self.headers(token),
        )
        self.assertEqual(response.status_code, 200, response.text)
        return response.json()

    def export(
        self,
        token: str,
        version_id: int,
        export_format: str = "docx",
        template_key: str = "professional",
        resume: dict | None = None,
        language: str = "en",
    ):
        if resume is None:
            resume = self.preview(token, version_id)["resume"]
        return self.client.post(
            f"/career/resume-versions/{version_id}/exports",
            headers=self.headers(token, True),
            json={
                "format": export_format,
                "template_key": template_key,
                "paper_size": "a4",
                "language": language,
                "resume": resume,
            },
        )

    def download(self, token: str, export_id: int):
        return self.client.get(
            f"/career/resume-exports/{export_id}/download",
            headers=self.headers(token),
        )

    def test_unauthenticated_user_cannot_preview_or_export(self) -> None:
        self.assertEqual(self.client.get("/career/resume-versions/1/preview").status_code, 401)
        self.assertEqual(
            self.client.post(
                "/career/resume-versions/1/exports", json={"format": "docx"}
            ).status_code,
            401,
        )

    def test_nonexistent_resume_version(self) -> None:
        token = self.register()["token"]
        self.assertEqual(
            self.client.get(
                "/career/resume-versions/9999/preview", headers=self.headers(token)
            ).status_code,
            404,
        )

    def test_preview_preserves_original_text_and_schema(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token)
        preview = self.preview(token, workspace["current_version"]["id"])
        self.assertEqual(
            preview["resume"]["original_text"], workspace["current_version"]["content"]
        )
        self.assertEqual(preview["resume"]["basics"]["name"], "Zheng Jieyao")
        self.assertEqual(preview["parse_status"], "structured")
        self.assertTrue(preview["resume"]["experience"])

    def test_unstructured_text_returns_needs_review(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token, "Zheng Jieyao\nBuilt reliable Python tools")
        preview = self.preview(token, workspace["current_version"]["id"])
        self.assertEqual(preview["parse_status"], "needs_review")
        self.assertTrue(preview["parse_warnings"])

    def test_professional_docx_generation_and_text(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token)
        response = self.export(token, workspace["current_version"]["id"])
        self.assertEqual(response.status_code, 201, response.text)
        downloaded = self.download(token, response.json()["id"])
        document = Document(io.BytesIO(downloaded.content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("Zheng Jieyao", text)
        self.assertIn("Built Python reporting workflows.", text)
        self.assertIn("EXPERIENCE", text)

    def test_minimal_ats_docx_has_no_tables(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token)
        response = self.export(
            token,
            workspace["current_version"]["id"],
            template_key="minimal_ats",
        )
        document = Document(io.BytesIO(self.download(token, response.json()["id"]).content))
        self.assertEqual(len(document.tables), 0)
        self.assertIn("SKILLS", "\n".join(p.text for p in document.paragraphs))

    def test_docx_headings_and_real_bullet_styles(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token)
        response = self.export(token, workspace["current_version"]["id"])
        document = Document(io.BytesIO(self.download(token, response.json()["id"]).content))
        self.assertTrue(any(p.style.name == "Resume Section" for p in document.paragraphs))
        bullets = [p for p in document.paragraphs if p.style.name == "List Bullet"]
        self.assertGreaterEqual(len(bullets), 3)

    def test_chinese_docx_content(self) -> None:
        chinese = """郑杰瑶
zheng@example.com

个人简介
专注于可靠的数据分析和业务决策支持。

工作经历
数据分析师
示例公司
2024 - 至今
- 使用 Python 建立自动化报表。

技能
Python，SQL，数据分析"""
        token = self.register()["token"]
        _, workspace = self.workspace(token, chinese)
        response = self.export(
            token,
            workspace["current_version"]["id"],
            language="zh",
        )
        downloaded = self.download(token, response.json()["id"])
        text = "\n".join(p.text for p in Document(io.BytesIO(downloaded.content)).paragraphs)
        self.assertIn("郑杰瑶", text)
        self.assertIn("使用 Python 建立自动化报表。", text)

    def test_english_docx_content(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token)
        response = self.export(token, workspace["current_version"]["id"])
        text = "\n".join(
            p.text
            for p in Document(io.BytesIO(self.download(token, response.json()["id"]).content)).paragraphs
        )
        self.assertIn("Data analyst focused on reliable decision support.", text)

    def test_empty_sections_are_not_rendered(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token)
        preview = self.preview(token, workspace["current_version"]["id"])
        preview["resume"]["awards"] = []
        response = self.export(token, workspace["current_version"]["id"], resume=preview["resume"])
        text = "\n".join(
            p.text
            for p in Document(io.BytesIO(self.download(token, response.json()["id"]).content)).paragraphs
        )
        self.assertNotIn("AWARDS", text)

    def test_manual_preview_edit_is_rendered_without_changing_version(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token)
        preview = self.preview(token, workspace["current_version"]["id"])
        preview["resume"]["basics"]["summary"] = "User-confirmed export summary."
        response = self.export(token, workspace["current_version"]["id"], resume=preview["resume"])
        text = "\n".join(
            paragraph.text
            for paragraph in Document(io.BytesIO(self.download(token, response.json()["id"]).content)).paragraphs
        )
        self.assertIn("User-confirmed export summary.", text)
        version = self.client.get(
            f"/career/resume-versions/{workspace['current_version']['id']}",
            headers=self.headers(token),
        ).json()
        self.assertNotIn("User-confirmed export summary.", version["content"])

    def test_source_content_mismatch_is_rejected(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token)
        preview = self.preview(token, workspace["current_version"]["id"])
        preview["resume"]["original_text"] = "different source"
        response = self.export(token, workspace["current_version"]["id"], resume=preview["resume"])
        self.assertEqual(response.status_code, 409)
        self.assertEqual(response.json()["detail"]["code"], "source_version_changed")

    def test_pdf_generation_and_expected_text(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token)
        response = self.export(token, workspace["current_version"]["id"], export_format="pdf")
        self.assertEqual(response.status_code, 201, response.text)
        downloaded = self.download(token, response.json()["id"])
        self.assertTrue(downloaded.content.startswith(b"%PDF"))
        text = "\n".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(downloaded.content)).pages)
        self.assertIn("Zheng Jieyao", text)
        self.assertIn("Built Python reporting workflows.", text)

    def test_chinese_pdf_content(self) -> None:
        chinese = """郑杰瑶
zheng@example.com

个人简介
专注于可靠的数据分析。

项目经历
智能求职工作台
个人项目
2025 - 至今
- 建立结构化简历导出流程。"""
        token = self.register()["token"]
        _, workspace = self.workspace(token, chinese)
        response = self.export(
            token,
            workspace["current_version"]["id"],
            export_format="pdf",
            language="zh",
        )
        self.assertEqual(response.status_code, 201, response.text)
        text = "\n".join(
            page.extract_text() or ""
            for page in PdfReader(io.BytesIO(self.download(token, response.json()["id"]).content)).pages
        )
        self.assertIn("郑杰瑶", text)
        self.assertIn("建立结构化简历导出流程。", text)

    def test_long_pdf_content_paginates(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token)
        preview = self.preview(token, workspace["current_version"]["id"])
        preview["resume"]["experience"] = [
            {
                "organization": f"Example Company {index}",
                "title": "Data Analyst",
                "location": "Shanghai",
                "start_date": "2020",
                "end_date": "2025",
                "bullet_points": [f"Synthetic achievement line {index}-{item} with Python and SQL." for item in range(20)],
            }
            for index in range(6)
        ]
        response = self.export(
            token,
            workspace["current_version"]["id"],
            export_format="pdf",
            resume=preview["resume"],
        )
        reader = PdfReader(io.BytesIO(self.download(token, response.json()["id"]).content))
        self.assertGreater(len(reader.pages), 1)

    def test_special_character_filename_is_safe(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token, company_name="ACME / 数据", job_title="Analyst: BI")
        preview = self.preview(token, workspace["current_version"]["id"])
        preview["resume"]["basics"]["name"] = "郑/杰瑶:*?"
        response = self.export(token, workspace["current_version"]["id"], resume=preview["resume"])
        filename = response.json()["filename"]
        self.assertNotRegex(filename, r"[\\/:*?\"<>|]")
        self.assertTrue(filename.endswith(".docx"))
        self.assertLessEqual(len(filename), 146)

    def test_directory_traversal_filename_is_removed(self) -> None:
        filename = ResumeExportService.safe_filename(
            "../../private", "../company", "..\\role", 3, "pdf"
        )
        self.assertNotIn("..", filename)
        self.assertNotIn("/", filename)
        self.assertNotIn("\\", filename)

    def test_pdf_renderer_unavailable_returns_clear_error(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token)
        renderer = self.client.app.state.resume_export_service.renderer
        with patch.object(
            renderer,
            "render_pdf",
            side_effect=DocumentRenderError("pdf_renderer_unavailable", "PDF unavailable"),
        ):
            response = self.export(token, workspace["current_version"]["id"], export_format="pdf")
        self.assertEqual(response.status_code, 503, response.text)
        self.assertEqual(response.json()["detail"]["code"], "pdf_renderer_unavailable")

    def test_export_failure_is_recorded_without_file(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token)
        renderer = self.client.app.state.resume_export_service.renderer
        with patch.object(renderer, "render", side_effect=RuntimeError("injected")):
            response = self.export(token, workspace["current_version"]["id"])
        self.assertEqual(response.status_code, 500)
        history = self.client.get(
            f"/career/resume-exports?version_id={workspace['current_version']['id']}",
            headers=self.headers(token),
        ).json()
        self.assertEqual(history[0]["status"], "failed")
        self.assertEqual(history[0]["error_code"], "export_generation_failed")
        self.assertEqual(list(self.export_dir.glob("*")), [])

    def test_export_record_creation_rolls_back(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token)
        repository = self.client.app.state.resume_export_service.repository
        with patch.object(repository, "_after_export_insert", side_effect=RuntimeError("injected")):
            with self.assertRaises(RuntimeError):
                self.export(token, workspace["current_version"]["id"])
        with sqlite3.connect(self.database_path) as connection:
            count = connection.execute("SELECT COUNT(*) FROM resume_exports").fetchone()[0]
        self.assertEqual(count, 0)

    def test_duplicate_exports_create_independent_records(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token)
        first = self.export(token, workspace["current_version"]["id"]).json()
        second = self.export(token, workspace["current_version"]["id"]).json()
        self.assertNotEqual(first["id"], second["id"])
        self.assertEqual(first["content_hash"], second["content_hash"])
        with sqlite3.connect(self.database_path) as connection:
            keys = [row[0] for row in connection.execute("SELECT object_key FROM resume_exports ORDER BY id")]
        self.assertEqual(len(set(keys)), 2)

    def test_export_history_persists_without_internal_path(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token)
        record = self.export(token, workspace["current_version"]["id"]).json()
        response = self.client.get(
            f"/career/resume-exports?version_id={workspace['current_version']['id']}",
            headers=self.headers(token),
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()[0]["id"], record["id"])
        serialized = response.text
        self.assertNotIn("object_key", serialized)
        self.assertNotIn(str(self.export_dir), serialized)

    def test_failed_export_cannot_be_downloaded(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token)
        renderer = self.client.app.state.resume_export_service.renderer
        with patch.object(renderer, "render", side_effect=RuntimeError("injected")):
            self.export(token, workspace["current_version"]["id"])
        with sqlite3.connect(self.database_path) as connection:
            export_id = connection.execute("SELECT id FROM resume_exports").fetchone()[0]
        response = self.download(token, export_id)
        self.assertEqual(response.status_code, 409)

    def test_download_content_type_and_disposition_docx(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token)
        record = self.export(token, workspace["current_version"]["id"]).json()
        response = self.download(token, record["id"])
        self.assertEqual(
            response.headers["content-type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertIn("attachment", response.headers["content-disposition"])
        self.assertIn(".docx", response.headers["content-disposition"])

    def test_download_content_type_pdf(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token)
        record = self.export(token, workspace["current_version"]["id"], export_format="pdf").json()
        response = self.download(token, record["id"])
        self.assertEqual(response.headers["content-type"], "application/pdf")

    def test_delete_export_removes_record_and_file(self) -> None:
        token = self.register()["token"]
        _, workspace = self.workspace(token)
        record = self.export(token, workspace["current_version"]["id"]).json()
        self.assertTrue(list(self.export_dir.glob("*.docx")))
        response = self.client.delete(
            f"/career/resume-exports/{record['id']}", headers=self.headers(token)
        )
        self.assertEqual(response.status_code, 204)
        self.assertEqual(list(self.export_dir.glob("*")), [])
        self.assertEqual(
            self.client.get(
                f"/career/resume-exports/{record['id']}", headers=self.headers(token)
            ).status_code,
            404,
        )

    def test_user_cannot_access_another_users_version_or_export(self) -> None:
        owner = self.register("export_owner")
        _, workspace = self.workspace(owner["token"])
        record = self.export(owner["token"], workspace["current_version"]["id"]).json()
        other = self.register("export_other")
        headers = self.headers(other["token"])
        self.assertEqual(
            self.client.get(
                f"/career/resume-versions/{workspace['current_version']['id']}/preview",
                headers=headers,
            ).status_code,
            404,
        )
        for path in (
            f"/career/resume-exports/{record['id']}",
            f"/career/resume-exports/{record['id']}/download",
        ):
            self.assertEqual(self.client.get(path, headers=headers).status_code, 404)
        self.assertEqual(
            self.client.delete(
                f"/career/resume-exports/{record['id']}", headers=headers
            ).status_code,
            404,
        )

    def test_frontend_exposes_real_export_workflow(self) -> None:
        page = (PROJECT_ROOT / "frontend" / "index.html").read_text(encoding="utf-8")
        script = (PROJECT_ROOT / "frontend" / "assets" / "js" / "app.js").read_text(encoding="utf-8")
        for element_id in (
            "optimizer-export-btn",
            "resume-export-workspace",
            "export-template-select",
            "export-format-select",
            "export-preview",
            "export-history",
        ):
            self.assertIn(f'id="{element_id}"', page)
        for function_name in (
            "openResumeExport",
            "createResumeExport",
            "downloadResumeExport",
            "deleteResumeExport",
        ):
            self.assertIn(f"function {function_name}", script)


if __name__ == "__main__":
    unittest.main()
